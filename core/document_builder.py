"""
document_builder.py — Generates a polished .docx from parsed BibEntry list.

Produces:
  • Cover page
  • Table of Contents placeholder
  • Main body (Title → Reference → Abstract per entry)
  • Processing Summary appendix
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Callable, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import BibEntry, FormatOptions, ParseResult
from .formatter import format_reference

# ── Colour palette ────────────────────────────────────────────────────────────
C_TITLE   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black navy
C_HEADING = RGBColor(0x2C, 0x3E, 0x50)   # dark slate
C_LABEL   = RGBColor(0x5A, 0x6A, 0x7A)   # medium grey
C_BODY    = RGBColor(0x1C, 0x1C, 0x1C)   # near-black
C_ACCENT  = RGBColor(0x3D, 0x6B, 0x8F)   # muted steel blue
C_LIGHT   = RGBColor(0x9A, 0xA0, 0xA6)   # light grey


# ── XML helpers ───────────────────────────────────────────────────────────────

def _page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run  = para.add_run()
    br   = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def _add_toc_field(doc: Document) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)

    run = para.add_run()
    for tag, attr, val in [
        ("w:fldChar",    "w:fldCharType", "begin"),
    ]:
        el = OxmlElement(tag)
        el.set(qn(attr), val)
        run._r.append(el)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    for val in ("separate", "end"):
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), val)
        run._r.append(fc)


def _para_border_bottom(para, color: str = "E0E0E0", sz: int = 4) -> None:
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _header_footer(doc: Document, header_text: str) -> None:
    for section in doc.sections:
        # Header
        hdr = section.header
        hdr.is_linked_to_previous = False
        hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run(header_text)
        r.font.name  = "Calibri"
        r.font.size  = Pt(8.5)
        r.font.color.rgb = C_LIGHT
        _para_border_bottom(hp, "D8D8D8", 3)

        # Footer — page numbers
        ftr = section.footer
        ftr.is_linked_to_previous = False
        fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _fld(run, instr: str) -> None:
            b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin"); run._r.append(b)
            i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = instr; run._r.append(i)
            e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end"); run._r.append(e)

        r0 = fp.add_run("— "); r0.font.name="Calibri"; r0.font.size=Pt(8.5); r0.font.color.rgb=C_LIGHT
        r1 = fp.add_run(); r1.font.name="Calibri"; r1.font.size=Pt(8.5); r1.font.color.rgb=C_LABEL; _fld(r1, " PAGE ")
        r2 = fp.add_run(" —"); r2.font.name="Calibri"; r2.font.size=Pt(8.5); r2.font.color.rgb=C_LIGHT


# ── Style helpers ─────────────────────────────────────────────────────────────

def _run(para, text: str, *, font: str, size: float,
         bold: bool = False, italic: bool = False,
         colour: Optional[RGBColor] = None) -> None:
    r = para.add_run(text)
    r.font.name  = font
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.italic = italic
    if colour:
        r.font.color.rgb = colour


def _spacing(para, before: float = 0, after: float = 0,
             line: float = 1.15) -> None:
    pf = para.paragraph_format
    pf.space_before      = Pt(before)
    pf.space_after       = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = line


def _apply_styles(doc: Document, opts: FormatOptions) -> None:
    """Override built-in heading/normal styles with user choices."""
    fn = opts.font_family

    # Normal
    ns = doc.styles["Normal"]
    ns.font.name = fn
    ns.font.size = Pt(opts.body_size_pt)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = fn; h1.font.size = Pt(opts.heading_size_pt + 4)
    h1.font.bold = True; h1.font.color.rgb = C_HEADING
    h1.paragraph_format.space_before = Pt(18); h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = fn; h2.font.size = Pt(opts.heading_size_pt)
    h2.font.bold = True; h2.font.color.rgb = C_HEADING
    h2.paragraph_format.space_before = Pt(12); h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = fn; h3.font.size = Pt(opts.heading_size_pt - 1)
    h3.font.bold = False; h3.font.italic = True
    h3.font.color.rgb = C_LABEL
    h3.paragraph_format.space_before = Pt(8); h3.paragraph_format.space_after = Pt(2)


# ── Section builders ──────────────────────────────────────────────────────────

def _build_cover(doc: Document, opts: FormatOptions,
                 included: int, total_raw: int) -> None:
    """Write a clean, centred cover page."""
    for _ in range(5):
        p = doc.add_paragraph(); _spacing(p, 0, 0)

    # Accent rule
    rule = doc.add_paragraph(); _spacing(rule, 0, 18)
    pPr  = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "18")
    top.set(qn("w:space"), "1"); top.set(qn("w:color"), "3D6B8F")
    pBdr.append(top); pPr.append(pBdr)

    # Main title
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _spacing(p, 0, 6)
    _run(p, opts.doc_title, font=opts.font_family, size=30,
         bold=True, colour=C_TITLE)

    # Subtitle
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _spacing(p, 0, 24)
    _run(p, opts.doc_subtitle, font=opts.font_family, size=14,
         italic=True, colour=C_LABEL)

    # Metadata block
    for label, value in [
        ("Date generated", datetime.now().strftime("%d %B %Y")),
        ("Entries included", f"{included:,}"),
        ("Source entries",   f"{total_raw:,}"),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _spacing(p, 0, 3)
        _run(p, f"{label}: ", font=opts.font_family, size=10, colour=C_LABEL)
        _run(p, value,        font=opts.font_family, size=10, bold=True, colour=C_HEADING)

    # Bottom rule
    for _ in range(3):
        pp = doc.add_paragraph(); _spacing(pp, 0, 0)
    rule2 = doc.add_paragraph(); _spacing(rule2, 12, 0)
    pPr2  = rule2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    bot   = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "BFBFBF")
    pBdr2.append(bot); pPr2.append(pBdr2)


def _build_toc_section(doc: Document, opts: FormatOptions) -> None:
    _page_break(doc)
    h = doc.add_heading("Table of Contents", level=1)
    _spacing(h, 0, 12)

    note = doc.add_paragraph()
    _spacing(note, 0, 16)
    _run(note,
         'Right-click this area in Microsoft Word and select "Update Field" '
         "to populate the table of contents with page numbers.",
         font=opts.font_family, size=9, italic=True, colour=C_LIGHT)

    _add_toc_field(doc)


def _build_article(doc: Document, idx: int, entry: BibEntry,
                   opts: FormatOptions) -> None:
    """Write a single article block: heading → reference → abstract."""
    fn = opts.font_family

    # Article number heading (H1 — captured in TOC)
    article_label = f"Article {idx}" if opts.number_entries else entry.title[:60]
    h1 = doc.add_heading(level=1)
    h1.clear()
    _spacing(h1, 22, 4)
    _run(h1, article_label, font=fn, size=opts.heading_size_pt + 2,
         bold=True, colour=C_HEADING)

    # ── Title block ───────────────────────────────────────────────────────
    lbl_t = doc.add_paragraph(); _spacing(lbl_t, 0, 1)
    _run(lbl_t, "TITLE", font=fn, size=7.5, bold=True, colour=C_ACCENT)

    p_t = doc.add_paragraph(); _spacing(p_t, 0, 8)
    _run(p_t, entry.title or "(No title)", font=fn,
         size=opts.body_size_pt + 1, bold=True, colour=C_BODY)

    # ── Reference block ───────────────────────────────────────────────────
    lbl_r = doc.add_paragraph(); _spacing(lbl_r, 0, 1)
    _run(lbl_r, "REFERENCE", font=fn, size=7.5, bold=True, colour=C_ACCENT)

    p_r = doc.add_paragraph(); _spacing(p_r, 0, 8, line=opts.line_spacing)
    _run(p_r, format_reference(entry), font=fn,
         size=opts.body_size_pt - 0.5, italic=False, colour=C_LABEL)

    # ── Abstract block ────────────────────────────────────────────────────
    lbl_a = doc.add_paragraph(); _spacing(lbl_a, 0, 1)
    _run(lbl_a, "ABSTRACT", font=fn, size=7.5, bold=True, colour=C_ACCENT)

    abs_text = entry.abstract if entry.has_abstract else "Not available."
    p_a = doc.add_paragraph(); _spacing(p_a, 0, 4, line=opts.line_spacing)
    if not entry.has_abstract:
        _run(p_a, abs_text, font=fn, size=opts.body_size_pt,
             italic=True, colour=C_LABEL)
    else:
        _run(p_a, abs_text, font=fn, size=opts.body_size_pt, colour=C_BODY)

    # Divider
    div = doc.add_paragraph(); _spacing(div, 10, 0)
    _para_border_bottom(div, "E8E8E8", 3)


def _build_summary(doc: Document, opts: FormatOptions,
                   result: ParseResult, included: int) -> None:
    _page_break(doc)
    fn = opts.font_family

    h = doc.add_heading("Processing Summary", level=1)
    _spacing(h, 0, 16)

    rows = [
        ("Total BibTeX entries detected",  result.total_raw),
        ("Entries successfully parsed",    len(result.entries)),
        ("Entries included in document",   included),
        ("Entries with abstracts",         len(result.with_abstract)),
        ("Entries missing abstracts",      len(result.without_abstract)),
        ("Malformed / skipped entries",    result.malformed_count),
        ("Duplicate titles detected",      len(result.duplicate_titles)),
    ]
    for label, value in rows:
        p = doc.add_paragraph(); _spacing(p, 0, 5)
        _run(p, f"{label}:  ", font=fn, size=10.5, bold=True, colour=C_LABEL)
        _run(p, f"{value:,}", font=fn, size=10.5, colour=C_BODY)

    if result.parse_errors:
        doc.add_paragraph()
        he = doc.add_heading("Parse Warnings", level=2)
        _spacing(he, 8, 6)
        for err in result.parse_errors[:20]:
            p = doc.add_paragraph(); _spacing(p, 0, 3)
            _run(p, err, font=fn, size=9.5, italic=True, colour=C_LABEL)

    if result.duplicate_titles:
        doc.add_paragraph()
        hd = doc.add_heading(f"Duplicate Titles ({len(result.duplicate_titles)})", level=2)
        _spacing(hd, 8, 6)
        for dt in result.duplicate_titles[:50]:
            p = doc.add_paragraph(); _spacing(p, 0, 3)
            _run(p, f"• {dt[:120]}", font=fn, size=9.5, colour=C_LABEL)

    # Generation stamp
    doc.add_paragraph()
    p = doc.add_paragraph(); _spacing(p, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p, f"Generated {datetime.now().strftime('%d %B %Y at %H:%M')}",
         font=fn, size=8.5, italic=True, colour=C_LIGHT)


# ── Public entry point ────────────────────────────────────────────────────────

def build_docx(
    entries: list[BibEntry],
    result: ParseResult,
    opts: FormatOptions,
    progress_callback: Optional[Callable[[float], None]] = None,
) -> bytes:
    """
    Build the Word document and return raw bytes ready for download.

    Args:
        entries:           Filtered and sorted list of entries to include.
        result:            Full ParseResult for summary statistics.
        opts:              User formatting options.
        progress_callback: Optional function(fraction: float) called during build.

    Returns:
        Raw .docx bytes.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    _apply_styles(doc, opts)
    _header_footer(doc, opts.doc_title)

    # Cover
    _build_cover(doc, opts, len(entries), result.total_raw)

    # TOC
    if opts.include_toc:
        _build_toc_section(doc, opts)

    # Articles section header
    _page_break(doc)
    h = doc.add_heading("Articles", level=1)

    total = len(entries)
    for idx, entry in enumerate(entries, start=1):
        if opts.page_breaks_between and idx > 1:
            _page_break(doc)
        _build_article(doc, idx, entry, opts)
        if progress_callback and idx % 50 == 0:
            progress_callback(idx / total)

    # Summary
    if opts.include_summary:
        _build_summary(doc, opts, result, len(entries))

    if progress_callback:
        progress_callback(1.0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
